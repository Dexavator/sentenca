import os
import json
import time
import asyncio
from datetime import datetime
from pathlib import Path
import uuid
from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file, jsonify
from flask_session import Session
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import PyPDF2
import google.generativeai as genai
import openai
from anthropic import Anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import secrets
from flask_wtf.csrf import CSRFProtect, generate_csrf
import tempfile
import shutil
import logging.handlers
from functools import wraps

# ====== DECORADORES ======
def handle_errors(f):
    @wraps(f)
    def wrapped(*args, **kwargs):
        try:
            return f(*args, **kwargs)
        except Exception as e:
            logger.error(f"Erro em {f.__name__}: {str(e)}")
            flash(f"Erro: {str(e)}", "danger")
            return redirect(url_for("index"))
    return wrapped

# ====== CARREGAR VARIÁVEIS DE AMBIENTE ======
load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")

if not GEMINI_API_KEY or not OPENAI_API_KEY or not ANTHROPIC_API_KEY:
    raise ValueError("Verifique se as chaves de API estão configuradas no .env")

# ====== CONFIGURAR APIs ======
genai.configure(api_key=GEMINI_API_KEY)
openai.api_key = OPENAI_API_KEY
anthropic = Anthropic(api_key=ANTHROPIC_API_KEY)

def create_app():
    app = Flask(__name__)
    
    # Configurações básicas
    app.secret_key = str(os.getenv('FLASK_SECRET_KEY'))
    
    # Usar /tmp para armazenamento temporário no Cloud Run
    temp_dir = tempfile.gettempdir()
    session_dir = os.path.join(temp_dir, 'flask_session')
    os.makedirs(session_dir, exist_ok=True)
    
    app.config.update(
        SESSION_TYPE='filesystem',
        SESSION_FILE_DIR=session_dir,
        SESSION_PERMANENT=False,
        SESSION_USE_SIGNER=True,
        SESSION_COOKIE_HTTPONLY=True,
        SESSION_COOKIE_SECURE=True,  # Alterado para True
        MAX_CONTENT_LENGTH=16 * 1024 * 1024,  # 16MB max-size
        WTF_CSRF_ENABLED=True,  # Alterado para True
        WTF_CSRF_CHECK_DEFAULT=True,
        WTF_CSRF_METHODS=['POST', 'PUT', 'PATCH', 'DELETE'],
        WTF_CSRF_SSL_STRICT=True,  # Alterado para True
        WTF_CSRF_TIME_LIMIT=None,
        WTF_CSRF_HEADERS=['X-CSRFToken']
    )
    
    # Inicializações
    csrf = CSRFProtect(app)
    Session(app)
    
    return app

app = create_app()

# ====== CONFIGURAÇÃO DO LOGGING ======
def setup_logging():
    """Configura o sistema de logging"""
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.DEBUG)

    # Handler para arquivo
    file_handler = logging.handlers.RotatingFileHandler(
        'app.log',
        maxBytes=1024 * 1024,  # 1MB
        backupCount=5
    )
    file_handler.setFormatter(logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    ))
    logger.addHandler(file_handler)

    # Handler para console
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(logging.Formatter(
        '%(levelname)s - %(message)s'
    ))
    logger.addHandler(console_handler)

    return logger

logger = setup_logging()

# ====== CLASSE PARA GERENCIAR HISTÓRICO DE SENTENÇAS ======
class SentencaHistory:
    def __init__(self):
        self.history_file = Path("sentencas_history.json")
        self.history = self._load_history()

    def _load_history(self):
        if self.history_file.exists():
            try:
                with open(self.history_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                return []
        return []

    def _save_history(self):
        with open(self.history_file, "w", encoding="utf-8") as f:
            json.dump(self.history, f, ensure_ascii=False, indent=4)

    def add_sentenca(self, sentenca, metadata):
        entry = {
            "id": str(uuid.uuid4()),
            "data": datetime.now().isoformat(),
            "sentenca": sentenca,
            "metadata": metadata
        }
        self.history.append(entry)
        self._save_history()

    def get_history(self):
        return self.history

# Inicializar o histórico
sentenca_history = SentencaHistory()

# ====== CLASSE PARA CACHE DE IA ======
class AICache:
    def __init__(self):
        self.cache_file = Path("ai_cache.json")
        self.cache = self._load_cache()

    def _load_cache(self):
        if self.cache_file.exists():
            try:
                with open(self.cache_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except:
                return {}
        return {}

    def _save_cache(self):
        with open(self.cache_file, "w", encoding="utf-8") as f:
            json.dump(self.cache, f, ensure_ascii=False, indent=4)

    def get(self, key):
        return self.cache.get(key)

    def set(self, key, value):
        self.cache[key] = value
        self._save_cache()

ai_cache = AICache()

# ====== FUNÇÕES PARA CHAMADAS ÀS IAs ======
def get_gemini_response(prompt, max_tentativas=5, base_espera=3):
    tentativa = 0
    while tentativa < max_tentativas:
        try:
            model = genai.GenerativeModel("gemini-2.0-flash-exp")
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            if "429" in str(e):
                tempo_espera = base_espera * (2 ** tentativa)
                logger.warning(f"Erro 429. Esperando {tempo_espera} segundos antes de tentar novamente.")
                time.sleep(tempo_espera)
                tentativa += 1
            else:
                logger.error(f"Erro ao chamar Gemini: {e}")
                raise
    raise Exception(f"Falha após {max_tentativas} tentativas. Erro 429.")

def get_openai_response(prompt, model_name="gpt-4o"):
    try:
        response = openai.ChatCompletion.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
        )
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"Erro ao chamar OpenAI: {e}")
        raise

def get_anthropic_response(prompt, model_name="claude-3-sonnet-20240229"):
    try:
        response = anthropic.messages.create(
            model=model_name,
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
        return response.content[0].text
    except Exception as e:
        logger.error(f"Erro ao chamar Anthropic: {e}")
        raise

async def get_ai_response(prompt, ai_choice, model_name=None):
    """Função assíncrona que usa cache."""
    cache_key = f"{ai_choice}:{model_name}:{prompt}"
    cached = ai_cache.get(cache_key)
    if cached:
        logger.debug(f"Cache hit para a chave: {cache_key}")
        return cached

    try:
        if ai_choice == "Gemini":
            # Não formatar o prompt para Gemini
            response = await asyncio.to_thread(get_gemini_response, prompt)
        elif ai_choice == "OpenAI":
            response = await asyncio.to_thread(get_openai_response, prompt, model_name)
        elif ai_choice == "Anthropic":
            # Formatar o prompt apenas para Anthropic
            formatted_prompt = f"\n\nHuman: {prompt}\n\nAssistant:"
            response = await asyncio.to_thread(get_anthropic_response, formatted_prompt, model_name)
        else:
            raise ValueError("IA não selecionada corretamente")

        ai_cache.set(cache_key, response)
        logger.debug(f"Cache hit para a chave: {cache_key}")
        return response

    except Exception as e:
        logger.error(f"Erro na chamada da IA: {e}")
        raise

# ====== ROTAS ======

@app.route("/", methods=["GET", "POST"])
def index():
    """Página inicial onde o usuário pode configurar a IA e carregar arquivos."""
    if request.method == "POST":
        # Verifica se é um upload de processo
        if "pdf_file" in request.files:
            file = request.files.get("pdf_file")
            if file and file.filename.lower().endswith(".pdf"):
                try:
                    # Salvar o arquivo usando nossa função auxiliar
                    final_filepath = save_uploaded_file(file, prefix="processo")
                    
                    # Extrair texto do PDF
                    text = ""
                    with open(final_filepath, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        for page in reader.pages:
                            extracted_text = page.extract_text()
                            if extracted_text:
                                text += extracted_text + "\n"

                    # Salvar o texto em um arquivo separado
                    unique_id = str(uuid.uuid4())
                    texto_path = os.path.join("static", f"{unique_id}_processo.txt")
                    with open(texto_path, "w", encoding="utf-8") as tf:
                        tf.write(text)
                    logger.debug(f"Texto do processo salvo em: {texto_path}")

                    # Guardar o caminho do arquivo na sessão
                    session["processo_text_path"] = texto_path
                    flash("Processo carregado com sucesso!", "success")
                    
                    # Não redireciona mais, retorna para a mesma página
                    return redirect(url_for("index"))
                    
                except Exception as e:
                    flash(f"Erro ao processar PDF: {str(e)}", "danger")
                    logger.error(f"Erro ao processar PDF: {e}")
                    return redirect(url_for("index"))

    # Se for GET ou se não for upload de processo
    return render_template("index.html")

@app.route("/set_ai_config", methods=["POST"])
def set_ai_config():
    """Salva as configurações de IA na sessão."""
    # Configurações para extração
    ai_choice_extracao = request.form.get("ai_choice_extracao")
    session["ai_choice_extracao"] = ai_choice_extracao
    if ai_choice_extracao == "OpenAI":
        session["openai_model_extracao"] = request.form.get("openai_model_extracao")
    elif ai_choice_extracao == "Anthropic":
        session["anthropic_model_extracao"] = request.form.get("anthropic_model_extracao")
    elif ai_choice_extracao == "Gemini":
        session["gemini_model_extracao"] = request.form.get("gemini_model_extracao")

    # Configurações para sentença
    ai_choice_sentenca = request.form.get("ai_choice_sentenca")
    session["ai_choice_sentenca"] = ai_choice_sentenca
    if ai_choice_sentenca == "OpenAI":
        session["openai_model_sentenca"] = request.form.get("openai_model_sentenca")
    elif ai_choice_sentenca == "Anthropic":
        session["anthropic_model_sentenca"] = request.form.get("anthropic_model_sentenca")
    elif ai_choice_sentenca == "Gemini":
        session["gemini_model_sentenca"] = request.form.get("gemini_model_sentenca")

    flash("Configurações de IA salvas com sucesso!", "success")
    return redirect(url_for("index"))

@app.route("/extrair_pedidos", methods=["GET", "POST"])
@handle_errors
def extrair_pedidos():
    """Extrai pedidos do processo usando IA."""
    # Verificar se existe processo carregado
    if "processo_text_path" not in session:
        flash("Nenhum processo carregado. Faça o upload primeiro.", "warning")
        return redirect(url_for("index"))

    if request.method == "POST":
        try:
            # Ler o texto do processo
            with open(session["processo_text_path"], "r", encoding="utf-8") as f:
                processo_text = f.read()

            # Configurar a IA
            ai_choice = session.get("ai_choice", "OpenAI")
            openai_model = session.get("openai_model", "gpt-4o")
            anthropic_model = session.get("anthropic_model", "claude-3-5-sonnet-latest")

            # Preparar o prompt
            prompt = (
                f"Identifique e liste todos os pedidos feitos pelo autor neste processo trabalhista, "
                f"cheque de novo para ter certeza que não falta pedido:\n\n"
                f"{processo_text}\n\n"
                f"Liste cada pedido em um item separado."
            )

            # Obter resposta da IA
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            pedidos_str = loop.run_until_complete(
                get_ai_response(
                    prompt,
                    ai_choice,
                    openai_model if ai_choice == "OpenAI" else anthropic_model
                )
            )
            loop.close()

            # Processar a resposta
            pedidos = [p.strip() for p in pedidos_str.splitlines() if p.strip()]
            
            # Guardar pedidos na sessão
            session["pedidos"] = pedidos
            flash("Pedidos extraídos com sucesso!", "success")
            return redirect(url_for("decidir_pedidos"))

        except Exception as e:
            flash(f"Erro ao extrair pedidos: {str(e)}", "danger")
            logger.error(f"Erro ao extrair pedidos: {e}")
            return redirect(url_for("extrair_pedidos"))

    # Se for GET, mostrar o formulário
    processo_text = ""
    if "processo_text_path" in session:
        with open(session["processo_text_path"], "r", encoding="utf-8") as f:
            processo_text = f.read()

    return render_template("extrair_pedidos.html", processo_text=processo_text)

@app.route("/decidir_pedidos", methods=["GET", "POST"])
@handle_errors
def decidir_pedidos():
    """Exibe os pedidos extraídos e permite ao usuário decidir sobre cada um."""
    # Verificar se existe processo carregado
    if "processo_text_path" not in session:
        flash("Nenhum processo carregado. Faça o upload primeiro.", "warning")
        return redirect(url_for("index"))

    # SEMPRE obter os pedidos da sessão, independentemente do método
    pedidos = session.get("pedidos", [])

    # Verificar se os pedidos foram extraídos
    if not pedidos:
        flash("Nenhum pedido foi extraído. Extraia os pedidos primeiro.", "warning")
        return redirect(url_for("extrair_pedidos"))

    if request.method == "POST":
        decisoes = []
        # Agora 'pedidos' está definido, mesmo no POST
        for i, pedido in enumerate(pedidos):
            decisao = request.form.get(f"decisao_{i}")
            incluir = request.form.get(f"incluir_{i}") == "on"
            artigos_selecionados = request.form.getlist(f"artigos_selecionados_{i}")
            decisoes.append({
                "pedido": pedido,
                "decisao": decisao,
                "incluir": incluir,
                "artigos": artigos_selecionados
            })
        session["decisoes"] = decisoes
        flash("Decisões salvas com sucesso!", "success")
        return redirect(url_for("gerar_sentenca"))

    return render_template("decidir_pedidos.html", pedidos=pedidos)

@app.route("/buscar_artigos", methods=["POST"])
def buscar_artigos():
    """Busca artigos relacionados ao pedido e decisão."""
    try:
        data = request.get_json()
        pedido = data.get('pedido')
        decisao = data.get('decisao')
        
        if not pedido or not decisao:
            return jsonify({
                'status': 'error',
                'message': 'Pedido e decisão são obrigatórios'
            }), 400

        # Aqui vai sua lógica para buscar artigos...
        artigos = [
            "Art. 7º da CF/88",
            "Art. 8º da CLT",
            "Súmula 331 do TST"
        ]

        return jsonify({
            'status': 'success',
            'artigos': artigos
        })

    except Exception as e:
        logger.error(f"Erro ao buscar artigos: {e}")
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route("/gerar_sentenca", methods=["GET", "POST"])
def gerar_sentenca():
    """Gera a sentença final (chamada à IA), usando a 'sentença modelo' e as decisões definidas."""
    logger.debug("Acessando a rota /gerar_sentenca")
    processo_text_path = session.get("processo_text_path", "")
    decisoes = session.get("decisoes", [])
    sentenca_modelo = session.get("sentenca_modelo", "")
    banco_decisoes = session.get("banco_decisoes", "")

    if request.method == "POST":
        # Montar prompts
        ai_choice = session.get("ai_choice", "OpenAI")
        openai_model = session.get("openai_model", "gpt-3.5-turbo")
        anthropic_model = session.get("anthropic_model", "claude-v1")
        logger.debug(f"Configurações de IA para gerar sentença: IA={ai_choice}, Modelo={openai_model if ai_choice == 'OpenAI' else anthropic_model}")

        # Instruções básicas
        instrucoes_pedidos = []
        for d in decisoes:
            if d["incluir"]:
                artigos_str = "\n    ".join(d["artigos"]) if d["artigos"] else "Nenhum artigo selecionado"
                instrucoes_pedidos.append(
                    f"- Pedido: {d['pedido']}\n"
                    f"  Decisão: {d['decisao']}\n"
                    f"  Artigos Selecionados:\n    {artigos_str}\n"
                )
        
        logger.debug(f"Instruções de pedidos: {instrucoes_pedidos}")

        if not instrucoes_pedidos:
            flash("Nenhum pedido está marcado para inclusão na sentença!", "warning")
            logger.warning("Nenhum pedido marcado para inclusão na sentença.")
            return redirect(url_for("decidir_pedidos"))

        # Se tiver banco de decisões, buscar trechos relevantes
        trechos_relevantes = ""
        if banco_decisoes:
            prompt_busca = f"""
            Analise o banco de decisões abaixo e extraia os trechos mais relevantes para cada pedido:

            PEDIDOS E DECISÕES:
            {''.join(instrucoes_pedidos)}

            BANCO DE DECISÕES:
            {banco_decisoes}

            INSTRUÇÕES:
            1. Para cada pedido, encontre trechos de fundamentação relacionados
            2. Priorize trechos que justifiquem a decisão de procedência/improcedência definida
            3. Inclua citações de leis e jurisprudência relevantes
            4. Mantenha apenas as partes mais pertinentes de cada trecho
            5. Organize os trechos por pedido

            Retorne apenas os trechos relevantes organizados por pedido.
            """
            try:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                trechos_relevantes = loop.run_until_complete(
                    get_ai_response(
                        prompt_busca,
                        ai_choice,
                        openai_model if ai_choice == "OpenAI" else anthropic_model
                    )
                )
                loop.close()
                logger.debug(f"Trechos relevantes extraídos: {trechos_relevantes}")
            except Exception as e:
                flash(f"Erro ao buscar trechos no banco de decisões: {str(e)}", "danger")
                logger.error(f"Erro ao buscar trechos no banco de decisões: {e}")

        # Montar o prompt para gerar a sentença
        if sentenca_modelo:
            # Prompt usando o modelo + trechos relevantes, se houver
            prompt_sentenca = f"""
            INSTRUÇÕES IMPORTANTES:
            Você deve gerar uma sentença trabalhista seguindo ESTRITAMENTE:
            1. As decisões de procedência/improcedência definidas para cada pedido
            2. O estilo, formatação e maneirismos EXATOS da sentença modelo fornecida
            3. Utilize os trechos relevantes do banco de decisões fornecidos abaixo

            PROCESSO:
            {processo_text_path}

            DECISÕES ESTABELECIDAS (ESTAS DECISÕES SÃO DEFINITIVAS E DEVEM SER SEGUIDAS À RISCA):
            {''.join(instrucoes_pedidos)}

            TRECHOS RELEVANTES DO BANCO DE DECISÕES:
            {trechos_relevantes}

            SENTENÇA MODELO (IMITE EXATAMENTE ESTE ESTILO):
            {sentenca_modelo}

            INSTRUÇÕES ESPECÍFICAS:
            1. Use EXATAMENTE a mesma estrutura da sentença modelo
            2. Imite o estilo de escrita, vocabulário e expressões específicas
            3. Incorpore os trechos relevantes do banco de decisões na fundamentação
            4. Adapte os trechos relevantes ao estilo da sentença modelo
            5. Mantenha a coerência entre os trechos utilizados e as decisões estabelecidas

            OBSERVAÇÕES:
            - Substitua sempre 'reclamante' e 'reclamada' e seus nomes por 'rte.' e 'rda.'
            - É CRUCIAL manter as decisões exatamente como informadas
            - Use os trechos relevantes para fortalecer a fundamentação

            IMPORTANTE: Esta sentença deve combinar o estilo do modelo com o conteúdo relevante do banco de decisões.
            """
        else:
            # Caso não haja sentença modelo, gerar com prompt genérico
            prompt_sentenca = f"""
            INSTRUÇÕES IMPORTANTES:
            Você deve gerar uma sentença trabalhista seguindo ESTRITAMENTE:
            1. As decisões de procedência/improcedência definidas para cada pedido
            2. O estilo, formatação e maneirismos EXATOS da sentença modelo fornecida

            PROCESSO:
            {processo_text_path}

            DECISÕES ESTABELECIDAS (ESTAS DECISÕES SÃO DEFINITIVAS E DEVEM SER SEGUIDAS À RISCA):
            {''.join(instrucoes_pedidos)}

            SENTENÇA MODELO (IMITE EXATAMENTE ESTE ESTILO):
            {sentenca_modelo}

            INSTRUÇÕES ESPECÍFICAS:
            1. Use EXATAMENTE a mesma estrutura da sentença modelo
            2. Imite o estilo de escrita, vocabulário e expressões específicas
            3. Mantenha o mesmo padrão de formatação (parágrafos, espaçamentos, etc.)
            4. Use as mesmas expressões e maneirismos do juiz que escreveu a sentença modelo
            5. Mantenha o mesmo nível de detalhamento e forma de argumentação
            6. Copie o estilo de citações de leis e jurisprudência
            7. Mantenha o mesmo padrão de transição entre as seções

            OBSERVAÇÕES:
            - Substitua sempre 'reclamante' e 'reclamada' e seus nomes por 'rte.' e 'rda.'
            - É CRUCIAL manter as decisões exatamente como informadas
            - A sentença deve parecer ter sido escrita pelo mesmo juiz que escreveu a sentença modelo

            IMPORTANTE: Esta sentença deve ser indistinguível em estilo e formato da sentença modelo fornecida.
            """

        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            sentenca_text = loop.run_until_complete(
                get_ai_response(
                    prompt_sentenca,
                    ai_choice,
                    openai_model if ai_choice == "OpenAI" else anthropic_model
                )
            )
            loop.close()
            logger.debug(f"Sentença gerada: {sentenca_text}")

            # Armazenar a sentença gerada na sessão
            session["sentenca_text"] = sentenca_text

            # Salvar no histórico
            metadata = {
                "ai_choice": ai_choice,
                "model": openai_model if ai_choice == "OpenAI" else anthropic_model,
                "pedidos": decisoes
            }
            sentenca_history.add_sentenca(sentenca_text, metadata)
            logger.debug("Sentença adicionada ao histórico")

            flash("Sentença gerada com sucesso!", "success")
            return redirect(url_for("sentenca_final"))

        except Exception as e:
            flash(f"Erro ao gerar sentença: {str(e)}", "danger")
            logger.error(f"Erro ao gerar sentença: {e}")
            return redirect(url_for("gerar_sentenca"))

    # Caso GET: exibir a página para o usuário poder escolher
    processo_text = ""
    try:
        with open(processo_text_path, "r", encoding="utf-8") as f:
            processo_text = f.read()
    except Exception as e:
        logger.error(f"Erro ao ler texto do processo: {e}")

    return render_template("gerar_sentenca.html", 
                           processo_text=processo_text,
                           decisoes=decisoes,
                           sentenca_modelo=sentenca_modelo,
                           banco_decisoes=banco_decisoes)

@app.route("/sentenca_final")
def sentenca_final():
    """Exibe a sentença final gerada para o usuário."""
    sentenca_text = session.get("sentenca_text", "")
    if not sentenca_text:
        flash("Nenhuma sentença gerada.", "warning")
        return redirect(url_for("gerar_sentenca"))
    return render_template("sentenca.html", 
                         sentenca_text=sentenca_text,
                         now=datetime.now())

# ====== EXPORTAR DOCX ======
@app.route("/export_docx")
def export_docx():
    sentenca_text = session.get("sentenca_text", "")
    if not sentenca_text:
        flash("Nenhuma sentença gerada.", "warning")
        logger.warning("Nenhuma sentença encontrada para exportação como DOCX.")
        return redirect(url_for("sentenca_final"))

    try:
        # Gerar DOCX em memória
        doc = Document()
        # Título
        title = doc.add_heading('SENTENÇA TRABALHISTA', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Remover marcações simples
        texto_limpo = sentenca_text.replace('**', '').replace('__', '').replace('```', '').replace('`', '').replace('#', '')
        paragrafos = texto_limpo.split('\n\n')

        for p in paragrafos:
            if p.strip():
                paragraph = doc.add_paragraph(p.strip())
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Salvar em objeto BytesIO para enviar
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Enviar como download
        return send_file(
            buffer,
            as_attachment=True,
            download_name="sentenca_trabalhista.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        flash(f"Erro ao exportar DOCX: {str(e)}", "danger")
        logger.error(f"Erro ao exportar DOCX: {e}")
        return redirect(url_for("sentenca_final"))

# ====== EXPORTAR PDF ======
@app.route("/export_pdf")
def export_pdf():
    sentenca_text = session.get("sentenca_text", "")
    if not sentenca_text:
        flash("Nenhuma sentença gerada.", "warning")
        logger.warning("Nenhuma sentença encontrada para exportação como PDF.")
        return redirect(url_for("sentenca_final"))

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        import io
        buffer = io.BytesIO()

        doc = SimpleDocTemplate(buffer, pagesize=letter)

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Sentenca',
            parent=styles['Normal'],
            fontSize=12,
            leading=14,
            alignment=4,  # Justificado
            spaceAfter=10
        ))

        Story = []
        title = Paragraph("<b>SENTENÇA TRABALHISTA</b>", styles["Heading1"])
        Story.append(title)
        Story.append(Spacer(1, 12))

        # Remover marcações simples
        texto_limpo = sentenca_text.replace('**', '').replace('__', '').replace('```', '').replace('`', '').replace('#', '')
        paragrafos = texto_limpo.split('\n\n')

        for p in paragrafos:
            if p.strip():
                p_escapado = p.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                p_escapado = p_escapado.replace('\n', '<br/>')
                paragraph = Paragraph(p_escapado, styles["Sentenca"])
                Story.append(paragraph)
                Story.append(Spacer(1, 12))

        doc.build(Story)
        buffer.seek(0)

        # Limpar arquivos temporários após o download
        processo_text_path = session.get("processo_text_path")
        if processo_text_path and os.path.exists(processo_text_path):
            os.remove(processo_text_path)
            session.pop("processo_text_path", None)  # Remover da sessão também

        sentenca_modelo_path = session.get("sentenca_modelo_path")
        if sentenca_modelo_path and os.path.exists(sentenca_modelo_path):
            os.remove(sentenca_modelo_path)
            session.pop("sentenca_modelo_path", None)

        banco_decisoes_path = session.get("banco_decisoes_path")
        if banco_decisoes_path and os.path.exists(banco_decisoes_path):
            os.remove(banco_decisoes_path)
            session.pop("banco_decisoes_path", None)

        # Limpar arquivos PDF antigos no diretório 'static'
        for filename in os.listdir("static"):
            if filename.endswith(".pdf"):
                filepath = os.path.join("static", filename)
                # Excluir arquivos com mais de 1 dia, por exemplo
                if os.path.getmtime(filepath) < time.time() - 24 * 60 * 60:
                    os.remove(filepath)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="sentenca_trabalhista.pdf",
            mimetype="application/pdf"
        )
    except Exception as e:
        flash(f"Erro ao exportar PDF: {str(e)}", "danger")
        logger.error(f"Erro ao exportar PDF: {e}")
        return redirect(url_for("sentenca_final"))

# ====== CARREGAR SENTENÇA MODELO ======
@app.route("/carregar_sentenca_modelo", methods=["POST"])
def carregar_sentenca_modelo():
    """Carrega uma sentença modelo de um arquivo."""
    file = request.files.get("modelo_file")
    if file:
        filename = secure_filename(file.filename)
        ext = os.path.splitext(filename)[1].lower()
        unique_id = str(uuid.uuid4())
        filepath = os.path.join("static", f"{unique_id}_{filename}")
        file.save(filepath)
        logger.debug(f"Arquivo de sentença modelo salvo em: {filepath}")

        texto_modelo = ""
        try:
            if ext == ".pdf":
                with open(filepath, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        extracted_text = page.extract_text()
                        if extracted_text:
                            texto_modelo += extracted_text + "\n"
            elif ext == ".docx":
                doc = Document(filepath)
                texto_modelo = '\n'.join([p.text for p in doc.paragraphs])
            else:
                with open(filepath, "r", encoding="utf-8") as f:
                    texto_modelo = f.read()

            # Salvar o texto da sentença modelo em um arquivo separado
            modelo_texto_path = os.path.join("static", f"{unique_id}_sentenca_modelo.txt")
            with open(modelo_texto_path, "w", encoding="utf-8") as tf:
                tf.write(texto_modelo)
            logger.debug(f"Texto da sentença modelo salvo em: {modelo_texto_path}")

            # Guardar o caminho do modelo na sessão
            session["sentenca_modelo_path"] = modelo_texto_path
            flash("Sentença modelo carregada com sucesso!", "success")
        except Exception as e:
            flash(f"Erro ao carregar sentença modelo: {str(e)}", "danger")
            logger.error(f"Erro ao carregar sentença modelo: {e}")

    return redirect(url_for("index"))

# ====== CARREGAR BANCO DE DECISÕES ======
@app.route("/carregar_banco_decisoes", methods=["POST"])
def carregar_banco_decisoes():
    """Carrega um arquivo com múltiplas decisões."""
    file = request.files.get("banco_file")
    if file:
        filename = secure_filename(file.filename)
        ext = os.path.splitext(filename)[1].lower()
        unique_id = str(uuid.uuid4())
        filepath = os.path.join("static", f"{unique_id}_{filename}")
        file.save(filepath)
        logger.debug(f"Arquivo de banco de decisões salvo em: {filepath}")

        texto_banco = ""
        try:
            if ext == ".pdf":
                with open(filepath, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        extracted_text = page.extract_text()
                        if extracted_text:
                            texto_banco += extracted_text + "\n"
            elif ext == ".docx":
                doc = Document(filepath)
                texto_banco = '\n'.join([p.text for p in doc.paragraphs])
            else:
                with open(filepath, "r", encoding="utf-8") as f:
                    texto_banco = f.read()

            # Salvar o texto do banco de decisões em um arquivo separado
            banco_texto_path = os.path.join("static", f"{unique_id}_banco_decisoes.txt")
            with open(banco_texto_path, "w", encoding="utf-8") as tf:
                tf.write(texto_banco)
            logger.debug(f"Texto do banco de decisões salvo em: {banco_texto_path}")

            # Guardar o caminho do banco de decisões na sessão
            session["banco_decisoes_path"] = banco_texto_path
            flash("Banco de decisões carregado com sucesso!", "success")
        except Exception as e:
            flash(f"Erro ao carregar banco de decisões: {str(e)}", "danger")
            logger.error(f"Erro ao carregar banco de decisões: {e}")

    return redirect(url_for("index"))

@app.route("/historico")
def historico():
    """Exibe o histórico de sentenças geradas."""
    history = sentenca_history.get_history()
    return render_template("historico.html", history=history)

def clean_temp_files():
    """Limpa arquivos temporários e antigos"""
    # Limpar arquivos da sessão
    for key in ['processo_text_path', 'sentenca_modelo_path', 'banco_decisoes_path']:
        path = session.get(key)
        if path and os.path.exists(path):
            try:
                os.remove(path)
                session.pop(key, None)
            except Exception as e:
                logger.error(f"Erro ao remover {path}: {e}")

    # Limpar TODOS os arquivos do /static
    try:
        for filename in os.listdir("static"):
            if any(filename.endswith(ext) for ext in ['.pdf', '.txt', '.docx']):
                filepath = os.path.join("static", filename)
                try:
                    os.remove(filepath)
                except Exception as e:
                    logger.error(f"Erro ao remover arquivo {filepath}: {e}")
    except Exception as e:
        logger.error(f"Erro ao limpar diretório static: {e}")

def save_uploaded_file(file, prefix=""):
    """Salva arquivo enviado de forma segura no diretório temporário"""
    try:
        filename = secure_filename(file.filename)
        unique_id = str(uuid.uuid4())
        final_name = f"{prefix}_{unique_id}_{filename}" if prefix else f"{unique_id}_{filename}"
        
        # Usar diretório temporário do sistema
        temp_dir = tempfile.gettempdir()
        static_dir = os.path.join(temp_dir, 'static')
        os.makedirs(static_dir, exist_ok=True)
        
        final_path = os.path.join(static_dir, final_name)
        
        # Criar um arquivo temporário e copiar o conteúdo
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        file.save(temp_file.name)
        temp_file.close()
        
        try:
            shutil.move(temp_file.name, final_path)
        except Exception as e:
            shutil.copy2(temp_file.name, final_path)
            os.unlink(temp_file.name)
        
        return final_path
    except Exception as e:
        if 'temp_file' in locals():
            try:
                os.unlink(temp_file.name)
            except:
                pass
        logger.error(f"Erro ao salvar arquivo: {e}")
        raise

@app.route("/limpar_cache", methods=["POST"])
def limpar_cache():
    """Limpa o cache de IA e arquivos temporários."""
    try:
        # Limpar cache de IA
        ai_cache.cache = {}
        ai_cache._save_cache()
        
        # Limpar arquivos temporários
        clean_temp_files()
        
        # Limpar arquivos da sessão
        session.clear()
        
        # Limpar diretório static
        for filename in os.listdir("static"):
            if any(filename.endswith(ext) for ext in ['.pdf', '.txt', '.docx']):
                filepath = os.path.join("static", filename)
                try:
                    os.remove(filepath)
                except Exception as e:
                    logger.error(f"Erro ao remover arquivo {filepath}: {e}")
        
        flash("Cache e arquivos temporários limpos com sucesso!", "success")
    except Exception as e:
        flash(f"Erro ao limpar cache: {str(e)}", "danger")
        logger.error(f"Erro ao limpar cache: {e}")
    
    return redirect(url_for("index"))

@app.route("/carregar_processo", methods=["GET", "POST"])
def carregar_processo():
    """Carrega o processo enviado pelo usuário."""
    if request.method == "GET":
        return redirect(url_for("index"))
        
    # Implemente aqui a lógica para carregar o processo
    file = request.files.get("processo_file")
    if file and file.filename.lower().endswith(".pdf"):
        try:
            # Salvar o arquivo usando a função auxiliar save_uploaded_file
            final_filepath = save_uploaded_file(file, prefix="processo")
            
            # Armazena o caminho do arquivo na sessão
            session["processo_text_path"] = final_filepath
            flash("Processo carregado com sucesso!", "success")
        except Exception as e:
            flash(f"Erro ao processar o processo: {str(e)}", "danger")
            logger.error(f"Erro ao processar o processo: {e}")
    else:
        flash("Por favor, selecione um arquivo PDF válido.", "warning")
        logger.warning("Arquivo inválido enviado para upload de processo")
    
    # Sempre retorna para a página inicial
    return redirect(url_for("index"))

@app.route("/carregar_todos_arquivos", methods=["POST"])
def carregar_todos_arquivos():
    """Carrega todos os arquivos de uma vez."""
    try:
        # Processar arquivo do processo (obrigatório)
        if "pdf_file" in request.files:
            processo_file = request.files["pdf_file"]
            if processo_file and processo_file.filename.lower().endswith(".pdf"):
                final_filepath = save_uploaded_file(processo_file, prefix="processo")
                
                # Extrair texto do PDF
                text = ""
                with open(final_filepath, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        extracted_text = page.extract_text()
                        if extracted_text:
                            text += extracted_text + "\n"

                # Salvar o texto em arquivo separado
                unique_id = str(uuid.uuid4())
                texto_path = os.path.join("static", f"{unique_id}_processo.txt")
                with open(texto_path, "w", encoding="utf-8") as tf:
                    tf.write(text)
                
                session["processo_text_path"] = texto_path
                flash("Processo carregado com sucesso!", "success")
            else:
                flash("Por favor, selecione um arquivo PDF válido para o processo.", "warning")
                return redirect(url_for("index"))

        # Processar sentença modelo (opcional)
        if "modelo_file" in request.files:
            modelo_file = request.files["modelo_file"]
            if modelo_file and modelo_file.filename:
                final_filepath = save_uploaded_file(modelo_file, prefix="modelo")
                session["sentenca_modelo_path"] = final_filepath
                flash("Sentença modelo carregada com sucesso!", "success")

        # Processar banco de decisões (opcional)
        if "banco_file" in request.files:
            banco_file = request.files["banco_file"]
            if banco_file and banco_file.filename:
                final_filepath = save_uploaded_file(banco_file, prefix="banco")
                session["banco_decisoes_path"] = final_filepath
                flash("Banco de decisões carregado com sucesso!", "success")

    except Exception as e:
        flash(f"Erro ao processar arquivos: {str(e)}", "danger")
        logger.error(f"Erro ao processar arquivos: {e}")

    return redirect(url_for("index"))

# ====== MAIN ======
if __name__ == "__main__":
    # Assegurar que o diretório para armazenar sessões existe
    session_dir = app.config['SESSION_FILE_DIR']
    os.makedirs(session_dir, exist_ok=True)

    # Assegurar que o arquivo de histórico existe
    if not os.path.exists(sentenca_history.history_file):
        sentenca_history._save_history()

    app.run(debug=True)
